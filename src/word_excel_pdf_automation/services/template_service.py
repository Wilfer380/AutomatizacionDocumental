from __future__ import annotations

import logging
import re
import shutil
from bisect import bisect_right
from pathlib import Path

from docx import Document

from ..config import PLACEHOLDER


logger = logging.getLogger(__name__)


class TemplateService:
    def validate_placeholder_present(self, template_path: Path, placeholder: str = PLACEHOLDER) -> bool:
        document = Document(template_path)
        return self._contains_placeholder(document, placeholder)

    def create_filled_copy(self, template_path: Path, destination_path: Path, replacement: str, placeholder: str = PLACEHOLDER) -> bool:
        destination_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(template_path, destination_path)
        document = Document(destination_path)
        replaced = self._replace_placeholder(document, placeholder, replacement)
        if not replaced:
            destination_path.unlink(missing_ok=True)
            return False
        document.save(destination_path)
        return True

    def _contains_placeholder(self, document: Document, placeholder: str) -> bool:
        for paragraph in self._iter_paragraphs(document):
            if placeholder in paragraph.text:
                return True
        return False

    def _replace_placeholder(self, document: Document, placeholder: str, replacement: str) -> bool:
        replaced = False
        for paragraph in self._iter_paragraphs(document):
            replaced |= self._replace_in_paragraph(paragraph, placeholder, replacement)
        return replaced

    def _iter_paragraphs(self, document: Document):
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            yield from self._iter_table_paragraphs(table)
        for section in document.sections:
            for paragraph in section.header.paragraphs:
                yield paragraph
            for table in section.header.tables:
                yield from self._iter_table_paragraphs(table)
            for paragraph in section.footer.paragraphs:
                yield paragraph
            for table in section.footer.tables:
                yield from self._iter_table_paragraphs(table)

    def _iter_table_paragraphs(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for inner_table in cell.tables:
                    yield from self._iter_table_paragraphs(inner_table)

    def _replace_in_paragraph(self, paragraph, placeholder: str, replacement: str) -> bool:
        if placeholder not in paragraph.text:
            return False
        if not paragraph.runs:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            return True

        run_texts = [run.text for run in paragraph.runs]
        full_text = "".join(run_texts)
        if placeholder not in full_text:
            return False

        boundaries = [0]
        for text in run_texts:
            boundaries.append(boundaries[-1] + len(text))

        matches = list(re.finditer(re.escape(placeholder), full_text))
        for match in reversed(matches):
            start = match.start()
            end = match.end() - 1
            start_run = bisect_right(boundaries, start) - 1
            end_run = bisect_right(boundaries, end) - 1
            start_offset = start - boundaries[start_run]
            end_offset = end - boundaries[end_run]

            if start_run == end_run:
                text = paragraph.runs[start_run].text
                paragraph.runs[start_run].text = text[:start_offset] + replacement + text[end_offset + 1 :]
                continue

            start_text = paragraph.runs[start_run].text
            end_text = paragraph.runs[end_run].text
            paragraph.runs[start_run].text = start_text[:start_offset] + replacement
            for index in range(start_run + 1, end_run):
                paragraph.runs[index].text = ""
            paragraph.runs[end_run].text = end_text[end_offset + 1 :]
        return True
